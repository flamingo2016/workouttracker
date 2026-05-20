from docx import Document
from docx.shared import Inches

risks = [
    {"Risk Description":"Incomplete inventory of Teradata assets due to limited legacy documentation","Probability":"High","Impact":"High","Mitigation":"Perform targeted discovery with DBAs; allocate SMEs for undocumented areas; use automated schema scanners","Owner":"Data Engineering Lead","Status":"Open","Response Plan":"Immediate discovery sprint; escalate missing docs to Sponsor; schedule extended profiling window"},
    {"Risk Description":"Data loss during bulk migration or ETL translation","Probability":"Medium","Impact":"High","Mitigation":"Implement checksum/hash reconciliation, row-count validation, and staged parallel runs; backup snapshots of source","Owner":"Migration Lead","Status":"Open","Response Plan":"Halt migration on discrepancies; restore from backups; run incremental re-sync"},
    {"Risk Description":"SQL dialect and stored procedure incompatibilities (Teradata->BigQuery)","Probability":"High","Impact":"High","Mitigation":"Inventory stored procs/macros; use automated conversion tools and SME review; build a conversion validation framework","Owner":"Data Platform Architect","Status":"Open","Response Plan":"Prioritise business-critical queries for manual conversion; allocate spike tasks for complex macros"},
    {"Risk Description":"Performance regressions post-migration impacting SLAs","Probability":"High","Impact":"High","Mitigation":"Establish performance baselines; run benchmarking; design slot reservations and BI Engine; query tuning cadence","Owner":"Performance Engineer","Status":"Open","Response Plan":"Rollback to Teradata for critical workloads; increase BigQuery resources temporarily; create optimization tickets"},
    {"Risk Description":"Regulatory non-compliance (SOX/HIPAA/CPNI handling)","Probability":"Medium","Impact":"High","Mitigation":"Engage Compliance early; document controls; enforce masking and access controls; audit logging via Cloud Logging","Owner":"Compliance Lead","Status":"Open","Response Plan":"Pause deployments for non-compliant domains; remediate controls and re-audit"},
    {"Risk Description":"CPNI exposure due to incomplete masking or misconfigured views","Probability":"Medium","Impact":"High","Mitigation":"GCP compliance team to implement masking for CPNI; review views and RBAC; test with privileged access workflows" ,"Owner":"Security Lead","Status":"Open","Response Plan":"Revoke problematic views; reimplement masking; notify affected stakeholders"},
    {"Risk Description":"Vendor SLA gaps for migration tooling or managed services","Probability":"Medium","Impact":"Medium","Mitigation":"Document required SLAs; maintain contingency vendors; negotiate interim terms; include penalties","Owner":"Procurement Lead","Status":"Open","Response Plan":"Escalate to Vendor Management; open change order or procure alternatives"},
    {"Risk Description":"Third-party BI tool compatibility and re-pointing risks","Probability":"Medium","Impact":"Medium","Mitigation":"Early testing with Tableau/Looker/Power BI; create connector templates; vendor engagement for drivers","Owner":"Integration Lead","Status":"Open","Response Plan":"Maintain dual connections during parallel run; provide fallback reports"},
    {"Risk Description":"Insufficient SME availability (20% allocated to opex)","Probability":"High","Impact":"Medium","Mitigation":"Resource levelling; hire/contract temporary specialists; protect key SME time in schedule","Owner":"Resource Manager","Status":"Open","Response Plan":"Reprioritise waves; request budget for contractors; escalate staffing to Sponsor"},
    {"Risk Description":"Fixed timeline pressure causing rushed validation","Probability":"High","Impact":"High","Mitigation":"Build formal quality gates; reserve buffer time in waves; obtain executive buy-in for scope-based trade-offs","Owner":"Project Manager","Status":"Open","Response Plan":"Defer non-critical waves; request timeline extension or reduce scope"},
    {"Risk Description":"Data quality issues discovered late during UAT","Probability":"Medium","Impact":"High","Mitigation":"Early profiling and data quality rules; create data repair playbooks; involve business SMEs","Owner":"Data Governance Lead","Status":"Open","Response Plan":"Implement corrective ETL rules; quarantine affected datasets; schedule remediation sprints"},
    {"Risk Description":"Security misconfigurations in GCP IAM or VPC exposing data","Probability":"Medium","Impact":"High","Mitigation":"Use least privilege, org policy constraints, VPC Service Controls, and periodic security reviews; automated scanning","Owner":"Cloud Security Lead","Status":"Open","Response Plan":"Revoke offending permissions; apply emergency org policies; conduct forensics"},
    {"Risk Description":"Cost overruns due to unexpected BigQuery or Storage usage","Probability":"Medium","Impact":"Medium","Mitigation":"Set budget alerts, cost governance, use reservations, and run cost modelling; enforce dataset lifecycle policies","Owner":"Finance Lead","Status":"Open","Response Plan":"Throttle heavy queries; enable cost controls; request contingency budget"},
    {"Risk Description":"Incomplete auditability or logging for regulatory evidence","Probability":"Low","Impact":"High","Mitigation":"Design logging requirements early; enable Cloud Logging and export logs to immutable storage; validate retention policies","Owner":"Compliance Lead","Status":"Open","Response Plan":"Enable missing logging, rebuild audit trails where possible, notify auditors"},
    {"Risk Description":"Parallel run complexity causing data divergence between Teradata and BigQuery","Probability":"Medium","Impact":"Medium","Mitigation":"Define clear dual-write or reconciliation strategy; frequent sync checkpoints; automated diff reports","Owner":"Migration Lead","Status":"Open","Response Plan":"Freeze cutover until divergence resolved; fallback to latest consistent snapshot"},
    {"Risk Description":"Decommissioning errors causing accidental loss of production Teradata data","Probability":"Low","Impact":"High","Mitigation":"Controlled decommission runbook, approvals, and hold periods; preserve backups and retention snapshots","Owner":"Operations Lead","Status":"Open","Response Plan":"Immediately restore from snapshot; delay decommission until validation complete"},
    {"Risk Description":"Legal or contract constraints with Teradata licensing during migration","Probability":"Low","Impact":"Medium","Mitigation":"Review contracts early with Legal; plan license transition; budget for buyouts if required","Owner":"Legal/Procurement","Status":"Open","Response Plan":"Negotiate transitional license terms; allocate legal support"},
    {"Risk Description":"Integration breaks for upstream systems still pointing to Teradata","Probability":"Medium","Impact":"Medium","Mitigation":"Inventory consumers; communicate cutover schedule; provide APIs/replication where needed","Owner":"Integration Lead","Status":"Open","Response Plan":"Repoint affected systems; maintain hybrid access during transition"},
    {"Risk Description":"Resistance to change and lack of business adoption post-migration","Probability":"Medium","Impact":"Medium","Mitigation":"Change management plan, training, and stakeholder engagement; early demos and pilot wins","Owner":"Change Lead","Status":"Open","Response Plan":"Run targeted training; collect feedback; provide business escalation routes"},
    {"Risk Description":"Backup/restore and DR strategy not validated in cloud","Probability":"Low","Impact":"High","Mitigation":"Implement and test DR runbooks; RTO/RPO validation; regional replication if required","Owner":"DR/Operations Lead","Status":"Open","Response Plan":"Trigger DR drills; failover to standby; remediate gaps"},
    {"Risk Description":"Dependency on single vendor/tool for SQL conversion creating lock-in risk","Probability":"Medium","Impact":"Medium","Mitigation":"Evaluate multiple tools; keep conversion artifacts portable; include source control for conversion scripts","Owner":"Tooling Lead","Status":"Open","Response Plan":"Switch to fallback tooling; roll back to manual conversion for critical items"},
]

document = Document()
document.add_heading('Risk Register: Teradata to GCP Migration (Telecom - Dallas, TX)', level=1)
document.add_paragraph('Estimated budget: $3.5M. Status: Planning phase. Use this register for workshop discussion and refinement.')

table = document.add_table(rows=1, cols=7)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
headers = ['Risk Description','Probability','Impact','Mitigation','Owner','Status','Response Plan']
for i, h in enumerate(headers):
    hdr_cells[i].text = h

for r in risks:
    row_cells = table.add_row().cells
    row_cells[0].text = r['Risk Description']
    row_cells[1].text = r['Probability']
    row_cells[2].text = r['Impact']
    row_cells[3].text = r['Mitigation']
    row_cells[4].text = r['Owner']
    row_cells[5].text = r['Status']
    row_cells[6].text = r['Response Plan']

# Save document
out_path = 'Risk_Register_Teradata_to_GCP.docx'
document.save(out_path)
print('Saved', out_path)
